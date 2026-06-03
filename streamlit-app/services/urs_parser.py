"""
URS文档解析器
支持格式：DOCX、PDF、TXT
"""
import os
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


def parse_docx(file_path: str) -> str:
    """解析DOCX文件，提取全部文本"""
    if Document is None:
        raise ImportError("python-docx not installed")
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables_text.append(" | ".join(cells))
    all_text = "\n".join(paragraphs)
    if tables_text:
        all_text += "\n\n[表格内容]\n" + "\n".join(tables_text)
    return all_text


def parse_pdf(file_path: str) -> str:
    """解析PDF文件，提取全部文本"""
    if pdfplumber is None:
        raise ImportError("pdfplumber not installed")
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
    return "\n\n".join(text_parts)


def parse_txt(file_path: str) -> str:
    """读取TXT文件内容"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_urs_file(file_path: str) -> str:
    """自动识别文件类型并解析"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return parse_docx(file_path)
    elif ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，支持：DOCX, PDF, TXT")
