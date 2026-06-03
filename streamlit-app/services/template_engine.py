"""
DOCX文档生成引擎 - 技术方案/配置清单/报价文件
支持默认模板和用户自定义模板（自动识别占位符）
"""
import io
import re
import os
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


# ==================== 占位符模式 ====================

PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')

# 预定义所有可用占位符
PLACEHOLDER_DEFINITIONS = {
    # 客户和项目信息
    "customer": "客户名称",
    "project_name": "项目名称",
    "date": "日期",
    "date_year": "年份",
    
    # 产品信息
    "product_name": "产品名称",
    "product_model": "产品型号",
    "product_spec": "产品规格说明",
    "product_price": "产品基价 (USD)",
    "total_price": "配置总价 (USD)",
    "total_price_cny": "配置总价 (人民币)",
    
    # 配置信息
    "config_summary": "配置摘要",
    "config_details": "配置明细表",
    "cavity_config": "腔体配置",
    "transfer_chamber": "传递舱",
    "material": "材质",
    "control_system": "控制系统",
    "validation": "验证文件",
    
    # URS分析
    "urs_summary": "URS需求摘要",
    "urs_params": "URS提取参数",
    "matched_keywords": "匹配关键词",
    "confidence": "匹配度",
    
    # 合规
    "compliance": "合规标准",
    "compliance_list": "合规标准列表",
    
    # 商务
    "price_range": "报价范围",
    "payment_terms": "付款条件",
    "delivery_time": "交货周期",
    "warranty": "质保期",
    "validity": "报价有效期",
    
    # 其他
    "signature": "签署栏",
}


# ==================== DOCX 生成函数 ====================

def _add_styled_paragraph(doc, text: str, style: str = "Normal", bold: bool = False,
                          size: int = 11, color: str = None, alignment=None, space_after: int = 6):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (1, 3, 5)])
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_table_from_data(doc, headers: List[str], rows: List[List[str]], col_widths: Optional[List[int]] = None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后间距


def generate_technical_proposal(data: Dict) -> bytes:
    """生成技术方案 DOCX"""
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    _add_styled_paragraph(doc, "技术方案", "Normal", bold=True, size=22,
                          color="#1677ff", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_styled_paragraph(doc, "Technical Proposal", "Normal", bold=False, size=11,
                          color="#888888", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # 基本信息
    customer = data.get("customer", "待确认客户")
    date_str = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    
    _add_styled_paragraph(doc, f"客户：{customer}", size=11, space_after=4)
    _add_styled_paragraph(doc, f"日期：{date_str}", size=11, space_after=4)
    _add_styled_paragraph(doc, f"文档编号：TP-{datetime.now().strftime('%Y%m%d')}-{hash(customer) % 10000:04d}", size=11, space_after=20)

    # 分隔线
    doc.add_paragraph("─" * 60)

    # ===== 1. 项目概述 =====
    _add_styled_paragraph(doc, "一、项目概述", bold=True, size=14, color="#1677ff", space_after=8)
    
    product_name = data.get("product_name", "")
    product_model = data.get("product_model", "")
    product_spec = data.get("product_spec", "")
    
    overview = (
        f"本方案针对 {customer} 的用户需求规范（URS），推荐采用 {product_name} ({product_model}) 型号产品。\n\n"
        f"{product_spec}\n\n"
        "该方案充分考虑了客户在工艺需求、洁净等级、合规要求等方面的URS输入，"
        "提供从设备交付到验证服务的完整解决方案。"
    )
    _add_styled_paragraph(doc, overview, size=11, space_after=12)

    # URS摘要
    urs_params = data.get("urs_params", {})
    if urs_params:
        _add_styled_paragraph(doc, "URS需求摘要：", bold=True, size=11, space_after=4)
        param_labels = {
            "chamber_size": "工作舱尺寸", "sterilization": "灭菌方式",
            "cleanliness": "洁净等级", "material": "材质要求",
            "compliance": "合规要求", "quantity": "需求数量", "control": "控制系统",
        }
        for key, label in param_labels.items():
            if key in urs_params and urs_params[key]:
                vals = "、".join(urs_params[key]) if isinstance(urs_params[key], list) else str(urs_params[key])
                _add_styled_paragraph(doc, f"  • {label}：{vals}", size=10, space_after=2)

    # ===== 2. 技术规格 =====
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, "二、技术规格与技术参数", bold=True, size=14, color="#1677ff", space_after=8)

    specs = [
        ("设备名称", product_name),
        ("设备型号", product_model),
        ("工作舱材质", data.get("material", "不锈钢304 (SUS304)")),
        ("表面处理", "内表面镜面抛光Ra≤0.4μm，外表面拉丝处理"),
        ("洁净等级", data.get("cleanliness", "ISO 5 (Class 100)")),
        ("密封方式", "充气密封/机械密封"),
        ("控制系统", data.get("control_system", "PLC + 触摸屏人机界面")),
        ("数据记录", "支持数据导出和审计追踪"),
    ]
    _add_table_from_data(doc, ["参数", "规格"], specs, col_widths=[5, 12])

    # ===== 3. 标准配置 =====
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, "三、标准配置清单", bold=True, size=14, color="#1677ff", space_after=8)

    config_items = data.get("config_details", [])
    if config_items:
        config_rows = [[i.get("item", ""), str(i.get("qty", 1)),
                        f"${i.get('price', 0):,.0f}" if i.get('price', 0) > 0 else "标配",
                        i.get("note", "")]
                       for i in config_items]
        _add_table_from_data(doc, ["配置项", "数量", "价格 (USD)", "备注"], config_rows, col_widths=[5, 2, 3, 7])

    # ===== 4. 合规与验证 =====
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, "四、合规与验证", bold=True, size=14, color="#1677ff", space_after=8)

    compliance_list = data.get("compliance", [])
    if compliance_list:
        for c in compliance_list:
            _add_styled_paragraph(doc, f"  ✅ {c}", size=11, space_after=2)

    _add_styled_paragraph(doc, "", size=6, space_after=2)
    _add_styled_paragraph(doc, "提供完整的验证文件包：", bold=True, size=11, space_after=4)
    for item in ["• 设计确认 (DQ)", "• 安装确认 (IQ)", "• 运行确认 (OQ)", "• 性能确认 (PQ) - 可选"]:
        _add_styled_paragraph(doc, f"  {item}", size=10, space_after=2)

    # ===== 5. 可选配置 =====
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, "五、可选配置项", bold=True, size=14, color="#1677ff", space_after=8)

    options_rows = [
        ["材质升级316L", "耐腐蚀不锈钢", "+$8,000 ~ $15,000"],
        ["远程监控系统", "Web/APP远程查看", "+$12,000"],
        ["PQ验证服务", "性能确认全套", "+$15,000 ~ $25,000"],
        ["温湿度监测", "实时记录+报警", "+$5,000"],
        ["审计追踪", "21 CFR Part 11合规", "+$10,000"],
    ]
    _add_table_from_data(doc, ["可选配置", "说明", "参考价格 (USD)"], options_rows, col_widths=[5, 6, 6])

    # 页脚
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, f"本方案由售前管理系统自动生成 | {date_str}", size=9,
                          color="#999999", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_config_list_docx(data: Dict) -> bytes:
    """生成配置清单 DOCX"""
    doc = Document()

    _add_styled_paragraph(doc, "配置清单", bold=True, size=20,
                          color="#1677ff", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    customer = data.get("customer", "")
    product_name = data.get("product_name", "")
    product_model = data.get("product_model", "")
    date_str = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    total_price = data.get("total_price", 0)

    _add_styled_paragraph(doc, f"客户：{customer}    产品：{product_name} ({product_model})    日期：{date_str}", size=10, space_after=16)

    config_items = data.get("config_details", [])
    if config_items:
        config_rows = []
        for i in config_items:
            config_rows.append([
                i.get("item", ""),
                str(i.get("qty", 1)),
                f"${i.get('price', 0):,.0f}" if i.get('price', 0) > 0 else "标准配置",
                i.get("note", ""),
            ])
        
        # 添加合计行
        total_base = sum(i.get("price", 0) for i in config_items)
        config_rows.append(["", "", f"${total_base:,.0f}", "合计 (USD)"])

        _add_table_from_data(doc, ["配置项", "数量", "价格 (USD)", "备注"], config_rows, col_widths=[5, 2, 3, 7])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_quotation_docx(data: Dict) -> bytes:
    """生成报价单 DOCX"""
    doc = Document()

    # 标题
    _add_styled_paragraph(doc, "报价单", "Normal", bold=True, size=22,
                          color="#1677ff", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_styled_paragraph(doc, "Quotation", "Normal", size=11,
                          color="#888888", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    customer = data.get("customer", "待确认客户")
    date_str = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    product_name = data.get("product_name", "")
    product_model = data.get("product_model", "")
    total_price = data.get("total_price", 0)
    validity = data.get("validity", "30天")

    # 报价信息
    info_rows = [
        ("客户名称", customer),
        ("报价日期", date_str),
        ("产品名称", f"{product_name} ({product_model})"),
        ("报价编号", f"Q-{datetime.now().strftime('%Y%m%d')}-{hash(customer + product_name) % 10000:04d}"),
        ("报价有效期", validity),
    ]
    _add_table_from_data(doc, ["项目", "内容"], info_rows, col_widths=[4, 13])

    # 价格表
    doc.add_paragraph()
    _add_styled_paragraph(doc, "价格明细：", bold=True, size=12, space_after=8)
    
    config_items = data.get("config_details", [])
    if config_items:
        price_rows = []
        for i in config_items:
            price_rows.append([
                i.get("item", ""),
                str(i.get("qty", 1)),
                f"${i.get('price', 0):,.0f}" if i.get('price', 0) > 0 else "标准配置",
            ])
        
        total_base = sum(i.get("price", 0) for i in config_items)
        price_rows.append(["", "", f"${total_base:,.0f}"])
        _add_table_from_data(doc, ["配置项", "数量", "价格 (USD)"], price_rows, col_widths=[8, 3, 6])

    # 商务条款
    doc.add_paragraph("─" * 60)
    _add_styled_paragraph(doc, "商务条款：", bold=True, size=12, color="#1677ff", space_after=8)
    terms = data.get("payment_terms", "预付30%，验收合格后70%")
    delivery = data.get("delivery_time", "签订合同后45-60天")
    warranty = data.get("warranty", "整机质保1年")
    
    _add_styled_paragraph(doc, f"• 付款条件：{terms}", size=11, space_after=4)
    _add_styled_paragraph(doc, f"• 交货周期：{delivery}", size=11, space_after=4)
    _add_styled_paragraph(doc, f"• 质保期限：{warranty}", size=11, space_after=4)
    _add_styled_paragraph(doc, f"• 报价有效期：{validity}", size=11, space_after=4)

    # 签署栏
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()
    _add_styled_paragraph(doc, "报价单位：", bold=True, size=11, space_after=4)
    _add_styled_paragraph(doc, f"日期：{date_str}", size=11, space_after=4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ==================== 占位符检测与模板填充 ====================

def analyze_template(docx_bytes: bytes) -> Dict[str, Any]:
    """
    分析上传的模板文件，检测占位符
    
    返回: {"placeholders": [占位符列表], "paragraph_count": 段落数}
    """
    doc = Document(io.BytesIO(docx_bytes))
    all_placeholders = set()

    for p in doc.paragraphs:
        found = PLACEHOLDER_PATTERN.findall(p.text)
        all_placeholders.update(found)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found = PLACEHOLDER_PATTERN.findall(cell.text)
                all_placeholders.update(found)

    # 用预定义定义补充说明
    placeholder_info = {}
    for ph in sorted(all_placeholders):
        placeholder_info[ph] = {
            "description": PLACEHOLDER_DEFINITIONS.get(ph, "自定义占位符"),
            "defined": ph in PLACEHOLDER_DEFINITIONS,
        }

    return {
        "placeholders": sorted(all_placeholders),
        "placeholder_info": placeholder_info,
        "paragraph_count": len(doc.paragraphs),
    }


def fill_template(docx_bytes: bytes, data: Dict) -> bytes:
    """
    用数据填充模板中的占位符
    
    data格式: {"customer": "客户名", "product_name": "产品名", ...}
    """
    doc = Document(io.BytesIO(docx_bytes))

    def replace_in_text(text: str) -> str:
        def replacer(match):
            key = match.group(1)
            val = data.get(key)
            if val is None:
                return match.group(0)  # 保持未匹配占位符
            return str(val)
        return PLACEHOLDER_PATTERN.sub(replacer, text)

    # 替换段落文本
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = replace_in_text(run.text)

    # 替换表格文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = replace_in_text(run.text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_all_documents(data: Dict, user_templates: Optional[Dict[str, bytes]] = None) -> Dict[str, bytes]:
    """
    生成全套文档
    
    返回: {"技术方案.docx": bytes, "配置清单.docx": bytes, "报价单.docx": bytes}
    """
    result = {}

    # 使用用户自定义模板或默认模板
    if user_templates and "技术方案" in user_templates:
        result["技术方案.docx"] = fill_template(user_templates["技术方案"], data)
    else:
        result["技术方案.docx"] = generate_technical_proposal(data)

    if user_templates and "配置清单" in user_templates:
        result["配置清单.docx"] = fill_template(user_templates["配置清单"], data)
    else:
        result["配置清单.docx"] = generate_config_list_docx(data)

    if user_templates and "报价单" in user_templates:
        result["报价单.docx"] = fill_template(user_templates["报价单"], data)
    else:
        result["报价单.docx"] = generate_quotation_docx(data)

    return result
